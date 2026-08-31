import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = docx.Document()

    # Set page margins (0.7 in)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Base styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(9.5)
    normal_style.font.color.rgb = RGBColor(24, 24, 27)

    # Helper function for cell background color
    def set_cell_background(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Helper function for cell margins
    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # Header section
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("K2 PEST CONTROL")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(190, 35, 32) # Brand Red
    title_p.paragraph_format.space_after = Pt(2)

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("Comprehensive Website Audit: Technical, SEO, Content & Local SEO Status")
    sub_run.font.size = Pt(12)
    sub_run.font.bold = True
    sub_run.font.color.rgb = RGBColor(82, 82, 91)
    sub_p.paragraph_format.space_after = Pt(8)

    # Metadata Box
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    cell_left = meta_table.cell(0, 0)
    cell_right = meta_table.cell(0, 1)
    cell_left.width = Inches(3.5)
    cell_right.width = Inches(3.5)
    
    set_cell_background(cell_left, "F8FAFC")
    set_cell_background(cell_right, "F8FAFC")
    set_cell_margins(cell_left, top=120, bottom=120, left=150, right=150)
    set_cell_margins(cell_right, top=120, bottom=120, left=150, right=150)

    p_l = cell_left.paragraphs[0]
    p_l.add_run("Target Domain: ").bold = True
    p_l.add_run("https://www.k2pc.ca\n")
    p_l.add_run("Technology Stack: ").bold = True
    p_l.add_run("Next.js 15+ / React 19 / MySQL / Prisma")
    p_l.paragraph_format.space_after = Pt(0)

    p_r = cell_right.paragraphs[0]
    p_r.add_run("Report Date: ").bold = True
    p_r.add_run("August 2026\n")
    p_r.add_run("Audit Scope: ").bold = True
    p_r.add_run("Technical, SEO, Content & Local GTA Setup")
    p_r.paragraph_format.space_after = Pt(0)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Critical Alert Callout
    alert_table = doc.add_table(rows=1, cols=1)
    alert_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    alert_cell = alert_table.cell(0, 0)
    alert_cell.width = Inches(7.0)
    set_cell_background(alert_cell, "FEF2F2")
    set_cell_margins(alert_cell, top=140, bottom=140, left=200, right=200)
    
    ap = alert_cell.paragraphs[0]
    ar_title = ap.add_run("⚠️ CRITICAL ASSESSMENT: CONTENT WRITING IS NOT COMPLETE\n")
    ar_title.bold = True
    ar_title.font.size = Pt(10.5)
    ar_title.font.color.rgb = RGBColor(185, 28, 28)

    ar_body = ap.add_run(
        "While the technical architecture, dynamic schema engine, lead-capture forms, and responsive UI components "
        "are successfully built, content writing across the website is currently INCOMPLETE. The site relies on placeholder "
        "phone numbers (555 numbers), dummy Saskatchewan Google Map coordinates, 3 sample blog articles, and stock testimonials "
        "that must be replaced with genuine, localized copy before launching live advertising or organic search campaigns."
    )
    ar_body.font.size = Pt(9)
    ar_body.font.color.rgb = RGBColor(127, 29, 29)
    ap.paragraph_format.space_after = Pt(0)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Scorecard Table
    score_table = doc.add_table(rows=1, cols=4)
    score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    scorecards = [
        ("Technical Core", "88%", "App Router & DB Fallback", "16A34A", "F0FDF4"),
        ("Technical SEO", "90%", "Sitemap, Robots & Schema", "2563EB", "EFF6FF"),
        ("Local SEO (GTA)", "70%", "11 Regions (Needs NAP fix)", "D97706", "FFFBEB"),
        ("Content Writing", "35%", "⚠️ Incomplete / Placeholders", "DC2626", "FEF2F2"),
    ]
    for i, (title, score, note, text_color, bg_color) in enumerate(scorecards):
        cell = score_table.cell(0, i)
        cell.width = Inches(1.75)
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        r1 = p.add_run(f"{title}\n")
        r1.bold = True
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = RGBColor(71, 85, 105)

        r2 = p.add_run(f"{score}\n")
        r2.bold = True
        r2.font.size = Pt(16)
        r2.font.color.rgb = RGBColor.from_string(text_color)

        r3 = p.add_run(f"{note}")
        r3.font.size = Pt(7.5)
        r3.font.color.rgb = RGBColor(100, 116, 139)
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section Helper
    def add_section_header(title):
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(15, 23, 42)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    def add_sub_header(title):
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(51, 65, 85)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)

    def style_table_headers(table, col_widths, headers):
        hdr_row = table.rows[0]
        for idx, (header_text, width) in enumerate(zip(headers, col_widths)):
            cell = hdr_row.cells[idx]
            cell.width = Inches(width)
            set_cell_background(cell, "F1F5F9")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(header_text)
            r.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(30, 41, 59)
            p.paragraph_format.space_after = Pt(0)

    def add_table_row(table, col_widths, data, is_even=False):
        row = table.add_row()
        for idx, (val, width) in enumerate(zip(data, col_widths)):
            cell = row.cells[idx]
            cell.width = Inches(width)
            if is_even:
                set_cell_background(cell, "FAFAFA")
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            p.paragraph_format.space_after = Pt(0)

    # 1. Technical View
    add_section_header("1. Technical View: What's Done vs. What's Remaining")
    add_sub_header("✅ Completed Technical Modules")

    tech_done_table = doc.add_table(rows=1, cols=3)
    tech_done_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.8, 4.0, 1.2]
    style_table_headers(tech_done_table, widths, ["Module", "Implementation Details", "Status"])
    
    tech_done_data = [
        ("Core Stack", "Next.js 15+ App Router, React 19, TypeScript, Tailwind CSS with fast SSR / SSG routing.", "Completed"),
        ("Database & Hybrid Fallback", "Prisma ORM with MySQL + in-memory file fallback (content-db.ts) ensuring 100% uptime if DB is down.", "Completed"),
        ("Lead Intake & Validation", "Multi-step booking form with Zod schema validation, XSS sanitization, anti-bot honeypot, and IP rate limiting.", "Completed"),
        ("Automated Email Engine", "Dual email dispatch (Admin alert + Customer branded confirmation) via Resend / SMTP.", "Completed"),
        ("Admin Portal", "JWT/Session authenticated dashboard for managing leads, services, blog articles, and business profile.", "Completed"),
        ("AI Chatbot Assistant", "Interactive assistant (/api/chat) with system prompt config, FAQs, and toggleable availability.", "Completed"),
        ("Font & Performance", "Zero layout shift via next/font (Space Grotesk, Inter, IBM Plex Mono) with preconnect resource hints.", "Completed"),
    ]
    for idx, row in enumerate(tech_done_data):
        add_table_row(tech_done_table, widths, [row[0], row[1], row[2]], idx % 2 == 1)

    add_sub_header("⏳ Remaining Technical Items")
    tech_rem_table = doc.add_table(rows=1, cols=3)
    tech_rem_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_rem = [1.8, 4.0, 1.2]
    style_table_headers(tech_rem_table, widths_rem, ["Task", "Description / Recommendation", "Priority"])

    tech_rem_data = [
        ("SMS Lead Dispatch", "Integrate Twilio / AWS SNS webhook to text on-call technicians immediately upon emergency submission.", "High"),
        ("Asset Host Migration", "Migrate pest catalog images from external Unsplash URLs to local optimized WebP/AVIF formats.", "Medium"),
        ("Security Headers & CSP", "Configure Content Security Policy, HSTS, and X-Frame-Options in next.config.ts.", "Medium"),
        ("Automated Testing", "Build Playwright E2E test suite for booking flows, API routes, and admin authentication.", "Low"),
    ]
    for idx, row in enumerate(tech_rem_data):
        add_table_row(tech_rem_table, widths_rem, [row[0], row[1], row[2]], idx % 2 == 1)

    doc.add_page_break()

    # 2. General SEO View
    add_section_header("2. General & Technical SEO: What's Done vs. What's Remaining")
    add_sub_header("✅ Completed SEO Features")

    seo_table = doc.add_table(rows=1, cols=3)
    seo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_table_headers(seo_table, widths, ["SEO Feature", "Implementation Details", "Status"])

    seo_data = [
        ("Dynamic XML Sitemap", "/sitemap.xml dynamically queries database for static routes, active services, and published blog posts.", "Completed"),
        ("Robots Directives", "/robots.txt allows public crawling while properly disallowing /admin/ and /api/ endpoints.", "Completed"),
        ("PestControlService Schema", "JSON-LD structured data in root layout with business hours, license number, accepted payments, and price range.", "Completed"),
        ("Service & Breadcrumb Schema", "Individual service pages embed dedicated Service and BreadcrumbList JSON-LD schemas.", "Completed"),
        ("FAQPage Schema", "Rich Snippet eligible FAQ structured data embedded on homepage and service pages for Google snippet expansions.", "Completed"),
        ("Meta & Social Tags", "Dynamic title templates, canonical tags, and OpenGraph/Twitter summary cards configured per page.", "Completed"),
        ("Google Verification & GA4", "Google Search Console verification token and GA4 tracking script (G-8EHS2WM33H) integrated.", "Completed"),
    ]
    for idx, row in enumerate(seo_data):
        add_table_row(seo_table, widths, [row[0], row[1], row[2]], idx % 2 == 1)

    add_sub_header("⏳ Remaining SEO Work")
    p_seo_rem = doc.add_paragraph()
    p_seo_rem.add_run("• ").bold = True
    p_seo_rem.add_run("Blog Article Schema: Add BlogPosting JSON-LD schema with author, datePublished, and dateModified to blog post pages.\n")
    p_seo_rem.add_run("• ").bold = True
    p_seo_rem.add_run("Internal Linking Architecture: Build contextual links between related service pages (e.g., cross-linking Rodent Control to Commercial IPM).\n")
    p_seo_rem.add_run("• ").bold = True
    p_seo_rem.add_run("301 Permanent Redirects: Add server-level permanent redirects in next.config.ts from alias routes (/privacy-policy -> /privacy, /terms-of-service -> /terms).\n")
    p_seo_rem.add_run("• ").bold = True
    p_seo_rem.add_run("Image Alt Text Audit: Ensure all CMS-uploaded pest images have keyword-optimized and descriptive alt text.\n")
    p_seo_rem.paragraph_format.space_after = Pt(6)

    # 3. Content Writing Status
    add_section_header("3. Content Writing Status: ⚠️ NOT COMPLETE (Gap Analysis)")
    
    content_table = doc.add_table(rows=1, cols=3)
    content_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_widths = [1.8, 1.4, 3.8]
    style_table_headers(content_table, c_widths, ["Section / Page", "Current State", "Missing Content & Required Work"])

    content_data = [
        ("Business Contact (NAP)", "Placeholder", "Phone is dummy (416) 555-0199; map coordinates point to Saskatoon instead of Toronto. Must provide official phone & address."),
        ("8 Core Service Pages", "Partial Draft", "Contains short descriptions and generic steps. Needs comprehensive Ontario preparation checklists, safety protocols, and warranty terms."),
        ("Blog & Resource Library", "3 Samples Only", "Only 3 sample posts exist. Needs 10–15 localized SEO pillar articles targeting GTA search queries (e.g., tenant bed bug rights, carpenter ants in North York)."),
        ("About Us Page", "Template Copy", "Uses placeholder team bios and stock avatars. Needs authentic company origin story, licensed exterminator profiles, and verified ministry credentials."),
        ("Commercial Services", "General Overview", "Needs dedicated industry breakdowns: Restaurant & Food Service Health Audits (DineSafe), Warehousing & Logistics, Multi-Unit Property Management."),
        ("Customer Testimonials", "Mock Reviews", "5 hardcoded placeholder reviews. Must be replaced with real, verified Google / Homestars reviews referencing specific GTA municipalities."),
    ]
    for idx, row in enumerate(content_data):
        add_table_row(content_table, c_widths, [row[0], row[1], row[2]], idx % 2 == 1)

    doc.add_page_break()

    # 4. Local SEO Analysis
    add_section_header("4. Local SEO (GTA) Analysis: Integrated vs. Missing")

    add_sub_header("📍 What is Currently Integrated in the Site")
    loc_int_table = doc.add_table(rows=1, cols=2)
    loc_int_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    loc_int_widths = [2.2, 4.8]
    style_table_headers(loc_int_table, loc_int_widths, ["Integrated Feature", "Current Implementation Details"])

    loc_int_data = [
        ("Ontario Geo Meta Headers", "Layout includes geo.region: 'CA-ON', geo.placename: 'Toronto', geo.position: '43.7142;-79.3364', and ICBM tags."),
        ("GTA Regional Scope", "Defined coverage for 11 GTA municipalities: Toronto, North York, Etobicoke, Scarborough, Mississauga, Brampton, Vaughan, Markham, Oakville, Richmond Hill, Burlington."),
        ("Local Service Area Badges", "ServiceAreaClient.tsx showcases region-specific response commitments (e.g., '2h Emergency Dispatch', 'Local Unit on Standby')."),
        ("PestControlService Schema", "Embeds complete Toronto PostalAddress (M3C 1H9), Ontario Applicator License (ON-849201-P), and local opening hours."),
        ("Interactive Map Module", "LocationMapWidget.tsx renders interactive Google Map directions link and neighborhood coverage selectors."),
        ("Direct Phone Call CTAs", "Mobile sticky header and emergency action banners configured with immediate tel: dialing protocols."),
    ]
    for idx, row in enumerate(loc_int_data):
        add_table_row(loc_int_table, loc_int_widths, [row[0], row[1]], idx % 2 == 1)

    add_sub_header("❌ What is Missing / Needs Improvement for Local SEO")
    loc_miss_table = doc.add_table(rows=1, cols=3)
    loc_miss_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    loc_miss_widths = [2.0, 3.8, 1.2]
    style_table_headers(loc_miss_table, loc_miss_widths, ["Missing Item", "Impact & Necessary Correction", "Priority"])

    loc_miss_data = [
        ("Correct NAP & Map Embed", "Critical Bug: Current embed coordinates in company.ts point to Saskatoon (-106.6834, 52.1504). Must update to Toronto Google Maps Place ID and real phone.", "Immediate"),
        ("Dedicated City Landing Pages", "All location links currently anchor to homepage widget. To rank across Peel, York, and Halton, build dedicated landing pages (/locations/mississauga, /locations/brampton, etc.).", "High"),
        ("Local Citation Links (sameAs)", "sameAs schema currently lacks links to authoritative Canadian directories: Google Business Profile, Homestars, BBB, YellowPages.ca, and TrustedPros.", "High"),
        ("GeoShape Multi-Polygon Schema", "Schema only has 1 coordinate point. Needs GeoCircle or multi-region ServiceArea polygons covering the 60km Greater Toronto service radius.", "Medium"),
        ("Municipal Bylaw Content", "Lack of city-specific FAQs regarding municipal property standards (e.g., Toronto Municipal Code Chapter 629 on pest control, Peel Region rodent control).", "Medium"),
        ("Live Google Reviews Sync", "AggregateRating in schema is hardcoded (480 reviews). Needs Google Place API integration or a direct link to the Google Review submission form.", "Medium"),
    ]
    for idx, row in enumerate(loc_miss_data):
        add_table_row(loc_miss_table, loc_miss_widths, [row[0], row[1], row[2]], idx % 2 == 1)

    # 5. Implementation Roadmap
    add_section_header("5. Actionable Implementation Roadmap")
    plan_table = doc.add_table(rows=1, cols=3)
    plan_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_widths = [1.5, 4.3, 1.2]
    style_table_headers(plan_table, p_widths, ["Phase", "Core Deliverables", "Timeline"])

    plan_data = [
        ("Phase 1: Critical Fixes", "• Fix Google Maps embed coordinates & verified business phone.\n• Add 301 redirects for legal route aliases.\n• Connect official Google Business Profile and social directory links.", "Days 1 – 3"),
        ("Phase 2: Content Writing", "• Write full copy for 8 service pages with prep guides & FAQs.\n• Publish 10 localized GTA blog articles.\n• Draft authentic About Us and Commercial sector pages.", "Week 1 – 2"),
        ("Phase 3: Local Expansion", "• Build dedicated City Landing Pages (Mississauga, Brampton, etc.).\n• Expand GeoShape polygon schema & embed Google Places review badge.\n• Integrate SMS lead notification for on-call technicians.", "Week 3 – 4"),
    ]
    for idx, row in enumerate(plan_data):
        add_table_row(plan_table, p_widths, [row[0], row[1], row[2]], idx % 2 == 1)

    # Footer note
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    p_foot = doc.add_paragraph()
    r_foot = p_foot.add_run("K2 Pest Control (k2pc.ca) • Generated by Antigravity AI Engineering • Confidential Website Audit")
    r_foot.font.size = Pt(8)
    r_foot.font.color.rgb = RGBColor(148, 163, 184)
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path = r"e:\Rapidtechpro\pestcontrol\k2pc_website_audit_report.docx"
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    create_report()
